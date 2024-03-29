# Author: Tharcisio Leone #
# Dataset: Filling Contract Automatically  #

## Join data from two different tables
# 0. Importing Libraries
# 1. Reading the data sets
# 2. Visualise the data set
# 3. Replace one single item and save it
# 4. Filling automatically all the items in a single contract
# 5. Filling automatically all the items in all contracts


# 0. Importing Libraries
import pandas as pd
import numpy as np
import openpyxl as opx
from docx import Document
from datetime import datetime


# 1. Reading the data set
document = Document('Contract.docx')
print(document)


# 2. Visualise the data set
for paragraph in document.paragraphs:
    print(paragraph.text)


# 3. Replace one single item and save it
for paragraph in document.paragraphs:
    paragraph.text = paragraph.text.replace('XXXX', 'Lira')

document.save('Contract_Lira.docx')


# 4. Filling automatically all the items in a single contract
document = Document('Contract.docx')

name = 'Lira'
item1 = 'Car'
item2 = 'Notebook'
item3 = 'Smartphone'

references = {
        'XXXX': name,
        'YYYY': item1,
        'ZZZZ': item2,
        'WWWW': item3,
        'DD': str(datetime.now().day),
        'MM': str(datetime.now().month),
        'AAAA': str(datetime.now().year),
    }

for paragraph in document.paragraphs:
    for code in references:
        value = references[code]
        paragraph.text = paragraph.text.replace(code, value)

document.save(f'Contract - {name}.docx')


# 5. Filling automatically ALL the items in ALL contracts
table = pd.read_excel('Information.xlsx')

for row in table.index:
    document = Document('Contract.docx')

    name = table.loc[row, 'Name']
    item1 = table.loc[row, 'Item1']
    item2 = table.loc[row, 'Item2']
    item3 = table.loc[row, 'Item3']

    references = {
        'XXXX': name,
        'YYYY': item1,
        'ZZZZ': item2,
        'WWWW': item3,
        'DD': str(datetime.now().day),
        'MM': str(datetime.now().month),
        'AAAA': str(datetime.now().year),
    }

    for paragraph in document.paragraphs:
        for code in references:
            value = references[code]
            paragraph.text = paragraph.text.replace(code, value)

    document.save(f'Contract - {name}.docx')

